from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
import os
class TemplateModel:
    """
    Represents a template model.

    Attributes:
        name (str): The name of the template.
        title (str): The title of the template.
        date (str): The date of the template.
        subject (str): The subject of the template.
        professor (str): The professor of the template.
        section (str): The section of the template.
        school (str): The school of the template.
    """

    def __init__(self, name: str, title: str, date, subject: str, professor: str = "", school: str = "", section: str = ""):
        self.name = name
        self.title = title
        self.date = date
        self.subject = subject
        self.professor = professor
        self.section = section
        self.school = school
        self.instructions = ""

        self.COLOR = RGBColor(0, 0, 0)
        self.FONT = "Cambria"

    def __str__(self):
        return f"{self.name} - {self.title} - {self.date} - {self.subject} - {self.professor} - {self.section}"

    def create_file(self, filename: str) -> str:
        """
        Creates a file with the template.

        Args:
            filename (str): The name of the file to save the template to.

        Returns:
            message (str): A message indicating the success of the operation.
        """
        doc = Document()
        doc.save(filename)

        return "File Created Successfully!"

    def add_instructions(self, instructions: str) -> None:
        """
        Adds instructions to the template.

        Args:
            instructions (str): The instructions to add to the template.
        """
        self.instructions = instructions

    def write_header(self, filename: str) -> str:
        """
        Creates a header for the template.

        Args:
            filename (str): The name of the file to save the header to.

        Returns:
            message (str): A message indicating the success of the operation.
        """
        if not os.path.exists(filename):
            return "File does not exist!"
        doc = Document(filename)

        # Add and Edit Section
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Add Header
        header = section.header

        # Edit Styles
        doc.styles["Heading 1"].font.size = Pt(20)
        doc.styles["Heading 1"].font.name = self.FONT
        doc.styles["Heading 1"].font.bold = True
        doc.styles["Heading 1"].font.color.rgb = self.COLOR

        doc.styles["Heading 2"].font.size = Pt(13)
        doc.styles["Heading 2"].font.name = self.FONT
        doc.styles["Heading 2"].font.bold = True
        doc.styles["Heading 2"].font.color.rgb = self.COLOR

        doc.styles["Body Text"].font.size = Pt(13)
        doc.styles["Body Text"].font.name = self.FONT
        doc.styles["Body Text"].font.bold = False
        doc.styles["Body Text"].font.color.rgb = self.COLOR

        # Adding Header Content
        header_content = [
            self.subject,
            self.title,
            f"{self.name}\t\t{self.date}"
        ]

        # SUBJECT
        header.paragraphs[0].text = header_content[0].upper()
        header.paragraphs[0].style = doc.styles["Heading 1"]
        header.paragraphs[0].paragraph_format.space_before = Pt(5)
        header.paragraphs[0].paragraph_format.space_after = Pt(0)

        # Title
        paragraph = header.add_paragraph(header_content[1], style = doc.styles["Heading 2"])
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(0)

        # Name and Date
        paragraph = header.add_paragraph(header_content[2], style = doc.styles["Body Text"])
        paragraph.paragraph_format.space_after = Pt(0)
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(1), alignment=WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Inches(6), alignment=WD_TAB_ALIGNMENT.RIGHT)

        # Professor and Section
        if self.professor or self.school or self.section:
            header_content.append(f"{self.professor}\t\t{self.school} - {self.section}")

            paragraph = header.add_paragraph(header_content[3], style = doc.styles["Body Text"])
            paragraph.paragraph_format.space_after = Pt(0)
            tab_stops = paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(1), alignment=WD_TAB_ALIGNMENT.LEFT)
            tab_stops.add_tab_stop(Inches(6), alignment=WD_TAB_ALIGNMENT.RIGHT)

        # Add Line
        header.add_paragraph("_"*90, style = doc.styles["Body Text"])

        doc.save(filename)
        return "Header Created Successfully!"


class HomeworkModel(TemplateModel):
    """
    Represents a homework template model.

    Attributes:
        name (str): The name of the homework.
        title (str): The title of the homework.
        date (str): The date of the homework.
        subject (str): The subject of the homework.
        questions (list[str]): The list of questions for the homework.
        professor (str): The professor of the template.
        section (str): The section of the template.
    """

    def __init__(self, name: str, title: str, date: str, subject: str, questions: list[str], professor: str = "", school: str = "", section: str = ""):
        super().__init__(name, title, date, subject, professor, school, section)
        self.questions = questions

    def write_questions(self, filename: str) -> str:
        # Check if filename exists
        if not os.path.exists(filename):
            return print("File does not exist!")

        doc = Document(filename)

        if self.instructions:
            paragraph = doc.add_paragraph(self.instructions, style = doc.styles["Body Text"])

        for question in self.questions:
            paragraph = doc.add_paragraph(question, style = doc.styles["List Number"])
            paragraph.paragraph_format.space_before = Pt(10)
            doc.add_paragraph("\n")

        doc.save(filename)

        print("Questions Written Successfully!")

class LetterModel(TemplateModel):
    """
    Represents a Template object.

    Args:
        name (str): The name of the writer.
        title (str): The title of the letter.
        date (str): The date of the letter.
        subject (str): The subject of the letter.
        recipient (str): The recipient of the letter.
        address (list[str]): The address of the letter. (Addres must be separated and list.split() by commas)
        professor (str, optional): The professor of the letter. Defaults to "".
        section (str, optional): The section of the letter. Defaults to "".
    """
    def __init__(self, name: str, title: str, date: str, subject: str, recipient: str, address: list[str], professor: str = "", school: str = "", section: str = ""):
        super().__init__(name, title, date, subject, professor, school, section)
        self.recipient = recipient
        self.address = address

    def write_letter(self, filename: str) -> str:
        # Check if filename exists
        if not os.path.exists(filename):
            return "File does not exist!"

        doc = Document(filename)

        if self.instructions:
            paragraph = doc.add_paragraph(self.instructions, style = doc.styles["Body Text"])

        # Add Address and Date
        paragraph = doc.add_paragraph(self.date, style = doc.styles["Body Text"])
        for line in self.address:
            paragraph = doc.add_paragraph(line, style = doc.styles["Body Text"])
            paragraph.paragraph_format.space_before = Pt(0)

        # Salutations
        paragraph = doc.add_paragraph(f"Dear {self.recipient}:\n", style = doc.styles["Body Text"])
        paragraph.paragraph_format.space_before = Pt(10)

        # Body
        paragraph = doc.add_paragraph("<Body of the letter goes here.>", style = doc.styles["Body Text"])
        paragraph.paragraph_format.space_before = Pt(10)

        # Closing
        paragraph = doc.add_paragraph("\nSincerely,\n", style = doc.styles["Body Text"])
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph = doc.add_paragraph(self.name, style = doc.styles["Body Text"])

        doc.save(filename)

        return "Recipient Written Successfully!"

class EssayModel(TemplateModel):
    """
    Represents an essay template model.

    Attributes:
        name (str): The name of the essay.
        title (str): The title of the essay.
        date (str): The date of the essay.
        subject (str): The subject of the essay.
        topic (str): The topic of the essay.
        professor (str): The professor of the template.
        section (str): The section of the template.
    """

    def __init__(self, name: str, title: str, date: str, subject: str, topic: str, professor: str = "", school: str = "", section: str = ""):
        super().__init__(name, title, date, subject, professor, school, section)
        self.topic = topic

    def write_essay(self, filename: str) -> str:
        # Check if filename exists
        if not os.path.exists(filename):
            return "File does not exist!"

        doc = Document(filename)

        if self.instructions:
            paragraph = doc.add_paragraph(self.instructions, style = doc.styles["Body Text"])

        # Add Title
        paragraph = doc.add_paragraph(self.title, style = doc.styles["Title"])
        # center the title
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(2)

        # Add Author
        paragraph = doc.add_paragraph(self.name, style = doc.styles["Heading 2"])
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add Intro, Body and Conclusion (justified)
        for cont in ["Intro", "Body", "Conclusion"]:
            paragraph = doc.add_paragraph(f"<{cont} Content goes here>", style = doc.styles["Body Text"])
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        doc.save(filename)

        return "Essay Written Successfully!"

class MathModel(TemplateModel):
    """
    Represents a math template model.

    Attributes:
        name (str): The name of the math model.
        title (str): The title of the math model.
        date (str): The date of the math model.
        subject (str): The subject of the math model.
        problems (list[str]): A list of problems in the math model.
        professor (str): The professor of the template.
        section (str): The section of the template.
    """

    def __init__(self, name: str, title: str, date: str, subject: str, problems: list[str], professor: str = "", school: str = "", section: str = ""):
        super().__init__(name, title, date, subject, professor, school, section)
        self.problems = problems

    def write_problems(self, filename: str) -> str:
        # check if filename exists
        if not os.path.exists(filename):
            return "File does not exist!"

        doc = Document(filename)

        if self.instructions:
            paragraph = doc.add_paragraph(self.instructions, style = doc.styles["Body Text"])

        for problem in self.problems:
            paragraph = doc.add_paragraph(problem, style = doc.styles["List Number"])
            paragraph.paragraph_format.space_before = Pt(10)
            doc.add_paragraph("\n")

        doc.save(filename)

        return "Problems Written Successfully!"

# if __name__ == "__main__":
#     template = TemplateModel("John Doe", "Example Homework", "July 06, 2024", "Mathematics", "Prof. Bill Joe", "Section 1")
#     template.create_file("example_template.docx")
#     template.write_header("example_template.docx")

    # homework = HomeworkModel("John Doe", "Example Homework", "July 06, 2024", "Mathematics", ["Question 1", "Question 2", "Question 3"], "Prof. Bill Joe", "Section 1")
    # homework.create_file("homework.docx")
    # homework.write_header("homework.docx")
    # homework.add_instructions("Answer the following questions:")
    # homework.write_questions("homework.docx")

    # letter = LetterModel("John Doe", "Example Letter", "July 06, 2024", "Mathematics", "Prof. Bill Joe", ["1234 Example St.", "City, State, Zip"], "Prof. Bill Joe", "Section 1")
    # letter.create_file("letter.docx")
    # letter.add_instructions("Write a letter to the professor.")
    # letter.write_letter("letter.docx")

    # essay = EssayModel("John Doe", "Example Essay", "July 06, 2024", "Mathematics", "Mathematics in the 21st Century", "Prof. Bill Joe", "Section 1")
    # essay.create_file("essay.docx")
    # essay.add_instructions("Write an essay on the following topic:")
    # essay.write_essay("essay.docx")

    # math = MathModel("John Doe", "Example Math", "July 06, 2024", "Mathematics", ["Problem 1", "Problem 2", "Problem 3"], "Prof. Bill Joe", "Section 1")
    # math.create_file("math.docx")
    # math.write_header("math.docx")
    # math.add_instructions("Solve the following problems:")
    # math.write_problems("math.docx")